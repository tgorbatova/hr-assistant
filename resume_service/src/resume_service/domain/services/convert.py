import enum
import tempfile
from io import BytesIO

import structlog
import textract
from docx import Document
from PyPDF2 import PdfReader

from resume_service.domain.exceptions.convert import UnsupportedTypeError
from resume_service.main.config import Settings

_logger: structlog.stdlib.BoundLogger = structlog.get_logger("inference.converter")


class FileType(enum.StrEnum):
    pdf = "pdf"
    docx = "docx"
    doc = "doc"


class ConverterService:
    def __init__(self, settings: Settings) -> None:
        self.settings = settings

    async def convert(self, file_data: dict) -> str:
        """Convert files to text."""
        filename = file_data["filename"].lower()
        content_type = file_data["content_type"]
        contents = file_data["content"]

        if (
            filename.endswith(".docx")
            or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            return await self.convert_docx(contents, filename)
        if filename.endswith(".doc") or content_type == "application/msword":
            return await self.convert_doc(contents, filename)
        if filename.endswith(".pdf") or content_type == "application/pdf":
            return await self.convert_pdf(contents, filename)
        raise UnsupportedTypeError(filename)

    @staticmethod
    async def convert_doc(contents: bytes, filename: str) -> str:
        """Конвертирует .doc файл в plain text с поддержкой RTF и многоязычного текста.

        :param contents:
        :param filename:
        :return: Строка с извлечённым текстом
        """
        try:
            # Сохраняем байты во временный файл с расширением .doc
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=True) as tmp_file:
                tmp_file.write(contents)
                tmp_file.flush()

                # Извлекаем текст с помощью textract
                text = textract.process(tmp_file.name, encoding="utf-8")
                return text.decode("utf-8")

        except Exception as e:
            _logger.error("Error converting DOC file %s: %s", filename, str(e))
            return ""

    @staticmethod
    async def convert_docx(contents: bytes, filename: str) -> str:
        """Convert .docx file to text."""
        try:
            _logger.info("Converting DOCX file %s", filename)
            doc = Document(BytesIO(contents))
            full_text = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            row_texts.append(text)
                    if row_texts:
                        full_text.append(" | ".join(row_texts))

            result = "\n".join(full_text)

            if result.strip():
                _logger.info("Successfully extracted text from DOCX file")
                return result.strip()

            _logger.warning("No text content found in DOCX file %s", filename)
            raise

        except Exception as e:
            _logger.error("Error converting DOCX file %s: %s", filename, str(e))
            raise

    @staticmethod
    async def convert_pdf(contents: bytes, filename: str) -> str:
        """Convert pdf to text."""
        try:
            _logger.info("Converting PDF file %s", filename)
            pdf_reader = PdfReader(BytesIO(contents))

            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            _logger.error("Error converting PDF file %s: %s", filename, str(e))
            raise
